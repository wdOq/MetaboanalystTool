import os
import time
import requests
import config
import pandas as pd
import metaboanalystbot
import reportgenerate
from docx import Document
from docx.shared import Inches
from PyPDF2 import PdfReader
from dotenv import load_dotenv
load_dotenv(dotenv_path="forllm.env")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_FOLDER = os.path.join(BASE_DIR, "static")
pdf_path = os.path.join(STATIC_FOLDER, "輸入文件檢查規則.pdf")

result = ""
def extract_pdf_rules(pdf_path):
    text = ""
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            text += page.extract_text()
    return text.strip()

def read_csv(csv_path):
    return pd.read_csv(csv_path)

def run_check_with_openrouter(rules, df, model="moonshotai/kimi-dev-72b:free"):
    csv_sample = df.to_csv(index=False)
    prompt = f"""
                You're a metabolomics assistant. Please check whether the CSV data complies with the following rules, excluding the fold change check.
                The rules are as follows: {rules}
                The original data is as follows: {csv_sample}
                Please identify whether there are any violations. Specify the problematic columns and explain the issues (examples may be listed). 
                If no errors are found, output the text: "Pass".If there are errors, output the text: "Fail" and provide details about the errors.
            """
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are an expert assistant responsible for checking biomedical CSV data."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2
    }

    response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        return f"Error：{response.status_code} - {response.text}"





def main(csv_path):
    rules = extract_pdf_rules(pdf_path)
    df = read_csv(csv_path)
    result = run_check_with_openrouter(rules, df)
    print("The result of the model check\n")
    print(result)
    if result.find('Pass') != -1:
        metaboanalystbot.crawler(csv_path)
        OUTPUT_FOLDER = os.path.join(BASE_DIR, "download")
        result_text = reportgenerate.analyze_image_with_llm(os.path.join(OUTPUT_FOLDER, "T-testChart.png"),
        "You are a reseacher specializing in metabolomics.Please analyze pictures I send you. It was a result from MetaboAnalyst which is a web-based platform dedicated for comprehensive metabolomics data analysis, interpretation and integration with other omics data. The picture include the analysis of fold change,T-test and Anova."); 
        print("LLM 回覆：", result_text)
        doc = Document()
        doc.add_heading('TargetedMetaCorePeakTable_wo_peakid_Analysis_1', level=0)
        doc.add_heading("MetaboAnalyst 網頁狀態回饋:", level=1)
        result = metaboanalystbot.get_submit_status_report()
        doc.add_paragraph(result) 
        doc.add_heading("LLM檢測分析之結果:", level=1)
        doc.add_paragraph(result_text)
        doc.add_heading("圖片結果顯示:", level=1)
        doc.add_picture(os.path.join(OUTPUT_FOLDER, "T-testChart.png"),width=Inches(5.0))
        doc.add_picture(os.path.join(OUTPUT_FOLDER, "FoldChangeChart.png"),width=Inches(5.0))
        doc.add_picture(os.path.join(OUTPUT_FOLDER, "ANOVAChart.png"),width=Inches(5.0))
        doc.add_picture(os.path.join(OUTPUT_FOLDER, "Adjusted_Label_Table_Normalization.png"),width=Inches(5.0))
        doc.add_picture(os.path.join(OUTPUT_FOLDER, "Adjusted_Label_Table_sample_normalization.png"),width=Inches(5.0))
        report_path = os.path.join(config.UPLOAD_FOLDER, "檢驗報告.docx")
        doc.save(report_path)
        return "report generatePass"
    else:
        print("The check failed."); 
        return result
